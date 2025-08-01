import os
import pytesseract
import google.generativeai as genai
import streamlit as st
from PIL import Image
from pdf2image import convert_from_bytes
from reportlab.pdfgen import canvas
from io import BytesIO
from dotenv import load_dotenv
import docx
from docx import Document

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Preformatted
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
import re
import html
# Add pandas and openpyxl for Excel handling
import pandas as pd
import openpyxl as xl
import logging
import logging.handlers
from datetime import datetime
import traceback
import sys
from functools import wraps
import time

import requests
st.set_page_config(
    page_title="Multi-Document OCR Assistant",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="expanded"
)

# Enhanced logging configuration
def setup_logging():
    """Set up comprehensive logging configuration"""
    
    # Create logs directory if it doesn't exist
    logs_dir = "logs"
    if not os.path.exists(logs_dir):
        os.makedirs(logs_dir)
    
    # Create custom formatter
    class CustomFormatter(logging.Formatter):
        """Custom formatter with colors for console output"""
        
        grey = "\x1b[38;21m"
        green = "\x1b[32;21m"
        yellow = "\x1b[33;21m"
        red = "\x1b[31;21m"
        bold_red = "\x1b[31;1m"
        reset = "\x1b[0m"
        
        FORMATS = {
            logging.DEBUG: grey + "%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s" + reset,
            logging.INFO: green + "%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s" + reset,
            logging.WARNING: yellow + "%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s" + reset,
            logging.ERROR: red + "%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s" + reset,
            logging.CRITICAL: bold_red + "%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s" + reset
        }
        
        def format(self, record):
            log_fmt = self.FORMATS.get(record.levelno)
            formatter = logging.Formatter(log_fmt, datefmt='%Y-%m-%d %H:%M:%S')
            return formatter.format(record)
    
    # Create root logger
    root_logger = logging.getLogger()
    root_logger.setLevel(logging.DEBUG)
    
    # Clear existing handlers
    root_logger.handlers.clear()
    
    # Console handler with custom formatter
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(CustomFormatter())
    
    # File handler for all logs
    file_handler = logging.handlers.RotatingFileHandler(
        filename=os.path.join(logs_dir, 'ocr_app.log'),
        maxBytes=10*1024*1024,  # 10MB
        backupCount=5,
        encoding='utf-8'
    )
    file_handler.setLevel(logging.DEBUG)
    file_formatter = logging.Formatter(
        '%(asctime)s [%(levelname)s] %(name)s:%(lineno)d - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(file_formatter)
    
    # Error file handler for errors and critical logs
    error_handler = logging.handlers.RotatingFileHandler(
        filename=os.path.join(logs_dir, 'ocr_app_errors.log'),
        maxBytes=5*1024*1024,  # 5MB
        backupCount=3,
        encoding='utf-8'
    )
    error_handler.setLevel(logging.ERROR)
    error_handler.setFormatter(file_formatter)
    
    # Add handlers to root logger
    root_logger.addHandler(console_handler)
    root_logger.addHandler(file_handler)
    root_logger.addHandler(error_handler)
    
    # Create specific loggers for different components
    loggers = {
        'main': logging.getLogger('ocr_app.main'),
        'file_processing': logging.getLogger('ocr_app.file_processing'),
        'gemini_api': logging.getLogger('ocr_app.gemini_api'),
        'document_generation': logging.getLogger('ocr_app.document_generation'),
        'ui': logging.getLogger('ocr_app.ui'),
        'performance': logging.getLogger('ocr_app.performance')
    }
    
    return loggers

# Performance monitoring decorator
def log_performance(logger):
    """Decorator to log function performance"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            start_time = time.time()
            function_name = f"{func.__module__}.{func.__name__}"
            
            logger.info(f"Starting {function_name}")
            try:
                result = func(*args, **kwargs)
                execution_time = time.time() - start_time
                logger.info(f"Completed {function_name} in {execution_time:.2f}s")
                return result
            except Exception as e:
                execution_time = time.time() - start_time
                logger.error(f"Error in {function_name} after {execution_time:.2f}s: {str(e)}")
                logger.error(f"Traceback: {traceback.format_exc()}")
                raise
        return wrapper
    return decorator

# Error handling decorator
def handle_errors(logger, default_return=None):
    """Decorator to handle and log errors gracefully"""
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                logger.error(f"Error in {func.__name__}: {str(e)}")
                logger.error(f"Full traceback: {traceback.format_exc()}")
                if default_return is not None:
                    logger.info(f"Returning default value: {default_return}")
                    return default_return
                raise
        return wrapper
    return decorator


# Initialize logging
loggers = setup_logging()
main_logger = loggers['main']
file_logger = loggers['file_processing']
gemini_logger = loggers['gemini_api']
doc_logger = loggers['document_generation']
ui_logger = loggers['ui']
perf_logger = loggers['performance']

# Log application startup
main_logger.info("="*50)
main_logger.info("OCR Multi-Document Application Starting")
main_logger.info(f"Python version: {sys.version}")
main_logger.info(f"Streamlit version: {st.__version__}")
main_logger.info("="*50)

# --- Local Ollama Configuration ---
OLLAMA_BASE_URL = "http://192.168.1.6:11434/"

def list_local_models():
    try:
        resp = requests.get(f"{OLLAMA_BASE_URL}/api/tags")
        resp.raise_for_status()
        return [model['name'] for model in resp.json().get("models", [])]
    except Exception as e:
        st.error(f"Error fetching local models: {str(e)}")
        return []

def ollama_generate(model: str, prompt: str):
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={"model": model, "prompt": prompt, "stream": False}
        )
        response.raise_for_status()
        return response.json().get("response", "No response from local model.")
    except Exception as e:
        return f"Error using local model '{model}': {str(e)}"

st.sidebar.subheader("AI Backend Selection")
if "backend_choice" not in st.session_state:
    st.session_state.backend_choice = "Gemini"

st.session_state.backend_choice = st.sidebar.radio(
    "Choose Model Backend",
    ["Gemini", "Local (Ollama)"],
    index=["Gemini", "Local (Ollama)"].index(st.session_state.backend_choice)
)

# --- Model selection for Ollama ---
if st.session_state.backend_choice == "Local (Ollama)":
    models = list_local_models()
    if models:
        if "local_model" not in st.session_state:
            st.session_state.local_model = models[0]

        selected_model = st.sidebar.selectbox("Choose Local Model", models, key="local_model")
    else:
        st.sidebar.warning("⚠️ No local models found at 192.168.1.7:11434")

@log_performance(perf_logger)
@handle_errors(gemini_logger, default_return="Error generating response")
def generate_chat_response(text, prompt, is_comparison=False):
    """
    Generate response using either Gemini or Ollama based on user selection
    """
    backend = st.session_state.get("backend_choice", "Gemini")
    selected_model = st.session_state.get("local_model", "llama2")

    full_prompt = (
        f"Compare and contrast the following documents:\n\n{text}\n\nPrompt: {prompt}"
        if is_comparison else f"{prompt}\n\n{text}"
    )

    if backend == "Local (Ollama)":
        gemini_logger.info(f"Using Local Ollama model: {selected_model}")
        try:
            response = requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={"model": selected_model, "prompt": full_prompt, "stream": False}
            )
            response.raise_for_status()
            return response.json().get("response", "No response from local model.")
        except Exception as e:
            error_msg = f"Error using local model '{selected_model}': {str(e)}"
            gemini_logger.error(error_msg)
            return error_msg
    else:
        gemini_logger.info("Using Gemini API")
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(full_prompt)
            return response.text if response else "No response from Gemini."
        except Exception as e:
            error_msg = f"Error with Gemini API: {str(e)}"
            gemini_logger.error(error_msg)
            return error_msg

@log_performance(perf_logger)
@handle_errors(gemini_logger, default_return="Error generating response")
def generate_chat_response(text, prompt, is_comparison=False):
    backend = st.session_state.get("backend_choice", "Gemini")
    selected_model = st.session_state.get("local_model", "llama2")

    full_prompt = (
        f"Compare and contrast the following documents:\n\n{text}\n\nPrompt: {prompt}"
        if is_comparison else f"{prompt}\n\n{text}"
    )

    if backend == "Local (Ollama)":
        gemini_logger.info(f"Using Local Ollama model: {selected_model}")
        try:
            response = requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={"model": selected_model, "prompt": full_prompt, "stream": False}
            )
            response.raise_for_status()
            return response.json().get("response", "No response from local model.")
        except Exception as e:
            error_msg = f"Error using local model '{selected_model}': {str(e)}"
            gemini_logger.error(error_msg)
            return error_msg
    else:
        gemini_logger.info("Using Gemini API")
        try:
            model = genai.GenerativeModel("gemini-2.0-flash")
            response = model.generate_content(full_prompt)
            return response.text if response else "No response from Gemini."
        except Exception as e:
            error_msg = f"Error with Gemini API: {str(e)}"
            gemini_logger.error(error_msg)
            return error_msg

# Load environment variables with logging
try:
    load_dotenv()
    main_logger.info("Environment variables loaded successfully")
    
    # Log environment info (without sensitive data)
    env_vars = ['GOOGLE_API_KEY']
    for var in env_vars:
        if os.getenv(var):
            main_logger.info(f"Environment variable {var}: [SET]")
        else:
            main_logger.warning(f"Environment variable {var}: [NOT SET]")
            
except Exception as e:
    main_logger.error(f"Failed to load environment variables: {str(e)}")

# Load Google API Key Securely
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    error_msg = "Google API Key is missing! Set it as an environment variable: GOOGLE_API_KEY"
    main_logger.critical(error_msg)
    st.error(error_msg)
    st.stop()

try:
    genai.configure(api_key=GOOGLE_API_KEY)
    gemini_logger.info("Gemini API configured successfully")
except Exception as e:
    gemini_logger.error(f"Failed to configure Gemini API: {str(e)}")
    st.error("Failed to configure Gemini API")
    st.stop()

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Initialize session state variables
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "transformed_prompt" not in st.session_state:
    st.session_state.transformed_prompt = ""

if "original_prompt" not in st.session_state:
    st.session_state.original_prompt = ""

if "needs_processing" not in st.session_state:
    st.session_state.needs_processing = False

if "documents" not in st.session_state:
    st.session_state.documents = {}  # {filename: extracted_text}

if "selected_docs" not in st.session_state:
    st.session_state.selected_docs = {}

tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"  # Update this path for your system
try:
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
        main_logger.info(f"Tesseract OCR path set to: {tesseract_path}")
    else:
        main_logger.warning(f"Tesseract OCR path does not exist: {tesseract_path}")
        main_logger.info("Attempting to use system Tesseract Installation")
except Exception as e:
    main_logger.error(f"Failed to set Tesseract OCR path: {str(e)}")

# Initialize session state variables
session_vars = [
    "chat_history", "transformed_prompt", "original_prompt",
    "needs_processing", "documents", "selected_docs"
]

for var in session_vars:
    if var not in st.session_state:
        if var == "chat_history":
            st.session_state.chat_history = []
        elif var in ["transformed_prompt", "original_prompt"]:
            st.session_state[var] = ""
        elif var == "needs_processing":
            st.session_state[var] = False
        elif var in ["documents", "selected_docs"]:
            st.session_state[var] = {}

        ui_logger.debug(f"Initialized session state variable: {var} with default value")

main_logger.info("Session state variables initialized successfully")

# Sidebar & Title
st.sidebar.write("👨 OCR Multi-Document Tool")
st.title("📝 Multi-Document OCR Chatbot")
st.sidebar.write("Upload images, PDFs, Excel files, or Word documents to extract text and get responses.")

@log_performance(perf_logger)
@handle_errors(file_logger, default_return="Error processing Excel file")
# Function to extract text from Excel files
def extract_text_from_excel(excel_file):
    file_logger.info(f"started processing Excel file: {getattr(excel_file, 'name', 'unknown')}")

    try:
        # Read the Excel file
        excel_data = pd.read_excel(excel_file, sheet_name=None)
        file_logger.info(f"Successfully read Excel file with {len(excel_data)} sheets")

        # Process each sheet
        all_text = []
        for sheet_name, df in excel_data.items():
            file_logger.debug(f"Processing sheet: {sheet_name} with {len(df)} rows and {len(df.columns)} columns")
            # Add sheet name as heading
            all_text.append(f"## Sheet: {sheet_name}")
            
            # Convert dataframe to string representation
            # Fill NaN values with empty string to avoid 'nan' in the output
            df_filled = df.fillna('')
            
            # Convert to markdown table format for better representation
            table_str = df_filled.to_markdown(index=False)
            all_text.append(table_str)
            # Add space between sheets
            all_text.append("\n")

            file_logger.debug(f"Processed sheet {sheet_name}: {len(table_str)} characters extracted")
        
        return "\n".join(all_text)
    except Exception as e:
        file_logger.info(f"Error ")
        return f"Error processing Excel file: {str(e)}"

@log_performance(perf_logger)
@handle_errors(file_logger, default_return="Error processing Word document")
# Function to extract text from Word documents
def extract_text_from_docx(docx_file):
    file_logger.info(f"Started processing Word document: {getattr(docx_file, 'name', 'unknown')}")
    try:
        doc = Document(docx_file)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        file_logger.info(f"Extracted {len(full_text)} paragraphs from Word document")
        
        # Extract text from tables
        for table in doc.tables:
            table_text = []
            # Create a header row with column numbers
            header_row = []
            for i in range(len(table.columns)):
                header_row.append(f"Column {i+1}")
            table_text.append(" | ".join(header_row))
            table_text.append("-" * (len(header_row) * 10))  # Add separator
            
            # Add table rows
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    # Get text from cell paragraphs
                    cell_text = " ".join([p.text for p in cell.paragraphs if p.text.strip()])
                    row_text.append(cell_text if cell_text else "")
                table_text.append(" | ".join(row_text))
            
            # Add the table as markdown
            full_text.append("\n".join(table_text))
        
        return "\n\n".join(full_text)
    except Exception as e:
        return f"Error processing Word document: {str(e)}"

# File uploader widget - expanded accepted file types
uploaded_files = st.sidebar.file_uploader("Upload Documents",
                                           type=["png", "jpg", "jpeg", "pdf", "xlsx", "xls", "docx", "doc"],
                                           accept_multiple_files=True)

# Clear chat and documents buttons
col1, col2 = st.sidebar.columns(2)
with col1:
    if st.button("🧹 Clear Chat"):
        st.session_state.chat_history = []
        st.success("Chat history cleared!")

def transform_prompt(input_prompt):
    if not input_prompt:
        return ""
    try:
        model = genai.GenerativeModel("gemini-2.0-flash")
        full_prompt = (
            "Refine the user prompt for an OCR-based document analysis system. Enhance clarity, precision, and extraction focus, preserving user intent. Output only the optimized prompt. provide shorter professional responses.\n\n"
            f"Original prompt: {input_prompt}"
        )
        response = model.generate_content(full_prompt)
        return response.text if response else input_prompt
    except Exception as e:
        print(f"Error transforming prompt: {str(e)}")
        return input_prompt

# Add sidebar section for manual prompt optimization
st.sidebar.subheader("Prompt Optimization")
sidebar_prompt = st.sidebar.text_area("Enter prompt to optimize:", key="sidebar_prompt")
if st.sidebar.button("Optimize Prompt"):
    if sidebar_prompt.strip():
        optimized = transform_prompt(sidebar_prompt)
        st.sidebar.success("Prompt optimized!")
        st.sidebar.write("Optimized prompt:")
        st.sidebar.info(optimized)
        # Copy to clipboard option
        if st.sidebar.button("Use this prompt"):
            st.session_state.original_prompt = optimized
            st.session_state.needs_processing = True
            st.rerun()

def extract_text_from_image(image):
    try:
        img = Image.open(image)
        return pytesseract.image_to_string(img)
    except Exception as e:
        return f"Error processing image: {str(e)}"

def extract_text_from_pdf(pdf_bytes):
    text = ""
    try:
        images = convert_from_bytes(pdf_bytes.read(), poppler_path=r"C:\Program Files\poppler-24.08.0\Library\bin")
        for img in images:
            text += pytesseract.image_to_string(img) + "\n"
        return text if text.strip() else "No text detected in the PDF."
    except Exception as e:
        return f"Error processing PDF: {str(e)}"
def _extract_tables_from_markdown(text: str) -> list:
    """
    Extract markdown tables from text and return tables and non-table content separately.
    
    Args:
        text: The markdown text to process
        
    Returns:
        A list of dictionaries, each with 'type' (either 'table' or 'text') and 'content'
    """
    # Pattern to match markdown tables - any line starting with | and containing | character
    table_pattern: str = r'((?:\n|^)\|[^\n]*\|[^\n]*(?:\n\|[^\n]*\|[^\n]*)*)'

    # Find all table sections
    tables: list = re.findall(table_pattern, text)

    # Replace tables with placeholders
    placeholder_text: str = text
    for i, table in enumerate(tables):
        placeholder_text = placeholder_text.replace(table, f"<<TABLE_{i}>>")

    # Split by placeholders to get non-table sections
    sections: list = re.split(r'<<TABLE_\d+>>', placeholder_text)

    # Reconstruct with tables and non-table sections properly separated
    current_idx: int = 0
    result: list = []
    for section in sections:
        if section:
            result.append({"type": "text", "content": section})
        if current_idx < len(tables):
            result.append({"type": "table", "content": tables[current_idx]})
            current_idx += 1

    return result

def _parse_markdown_table(table_markdown: str) -> list:
    """
    Parse a markdown table into rows and cells.
    
    Args:
        table_markdown: The markdown table text
        
    Returns:
        A list of rows, where each row is a list of cell values
    """
    lines: list = table_markdown.strip().split('\n')
    rows: list = []

    # Skip separator row (contains only |, -, :)
    for i, line in enumerate(lines):
        if line.strip() and i != 1:  # Skip the separator row
            # Extract cells from each row
            cells: list = line.split('|')
            # Remove empty first/last cells (due to leading/trailing |)
            if not cells[0].strip():
                cells = cells[1:]
            if not cells[-1].strip():
                cells = cells[:-1]
            # Clean up whitespace in cells
            cells = [cell.strip() for cell in cells]
            rows.append(cells)

    return rows

def generate_pdf(text):
    """
    Generate PDF with proper formatting, including tables, with robust error handling
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, 
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=72)
    
    styles = getSampleStyleSheet()
    # Add custom style for code blocks
    if 'CodeBlock' not in styles:
        styles.add(ParagraphStyle(name='CodeBlock',
                                 fontName='Courier',
                                 fontSize=9,
                                 leftIndent=36,
                                 rightIndent=36,
                                 spaceAfter=12,
                                 backColor=colors.lightgrey))
    
    # Define table style
    table_style = TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('PADDING', (0, 0), (-1, -1), 6),
    ])
    
    def clean_text_for_pdf(text_content):
        """Clean text content to avoid ReportLab parsing errors"""
        if not text_content:
            return ""
        
        # Convert to string first
        text_str = str(text_content)
        
        # Escape HTML entities first
        cleaned = html.escape(text_str)
        
        # Remove or escape problematic characters
        replacements = {
            # Remove null bytes and other control characters
            '\x00': '', '\x01': '', '\x02': '', '\x03': '', '\x04': '',
            '\x05': '', '\x06': '', '\x07': '', '\x08': '', '\x0b': '',
            '\x0c': '', '\x0e': '', '\x0f': '',
        }
        
        for old, new in replacements.items():
            cleaned = cleaned.replace(old, new)
        
        # Remove any remaining control characters (except \n, \r, \t)
        cleaned = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', cleaned)
        
        return cleaned
    
    def safe_add_paragraph(text_content, style):
        """Safely add a paragraph with error handling"""
        try:
            cleaned_text = clean_text_for_pdf(text_content)
            if cleaned_text.strip():
                return Paragraph(cleaned_text, style)
            else:
                return Spacer(1, 6)
        except Exception as e:
            # If paragraph creation fails, use Preformatted as fallback
            try:
                safe_text = str(text_content).replace('<', '&lt;').replace('>', '&gt;').replace('&', '&amp;')
                return Preformatted(safe_text, styles['Code'])
            except:
                # Last resort: return a simple error message
                return Paragraph("Error: Could not process this text content.", styles['Normal'])
    
    # Extract tables and other content
    try:
        content_parts = _extract_tables_from_markdown(text)
    except Exception as e:
        # If markdown parsing fails, treat entire content as text
        content_parts = [{"type": "text", "content": str(text)}]
    
    # Build document elements
    elements = []
    
    for part in content_parts:
        if part["type"] == "text":
            # Process non-table text
            try:
                paragraphs = part["content"].split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        if para.startswith('```') and para.endswith('```'):
                            # Code block - use Preformatted to avoid parsing issues
                            code = para.strip('`').strip()
                            try:
                                clean_code = clean_text_for_pdf(code)
                                elements.append(Preformatted(clean_code, styles['Code']))
                            except:
                                # Fallback for problematic code blocks
                                elements.append(Paragraph("Code block (could not display)", styles['Normal']))
                        else:
                            # Regular paragraph - use safe method
                            paragraph_element = safe_add_paragraph(para, styles['Normal'])
                            elements.append(paragraph_element)
                            elements.append(Spacer(1, 6))
            except Exception as e:
                # If text processing fails, add as simple paragraph
                elements.append(safe_add_paragraph(str(part["content"]), styles['Normal']))
        
        elif part["type"] == "table":
            # Process table
            try:
                table_data = _parse_markdown_table(part["content"])
                if table_data and len(table_data) > 0:
                    # Clean table data
                    cleaned_table_data = []
                    for row in table_data:
                        cleaned_row = []
                        for cell in row:
                            # Clean each cell content
                            cleaned_cell = clean_text_for_pdf(str(cell)) if cell else ""
                            cleaned_row.append(cleaned_cell)
                        cleaned_table_data.append(cleaned_row)
                    
                    # Create Table object
                    table = Table(cleaned_table_data)
                    table.setStyle(table_style)
                    elements.append(table)
                    elements.append(Spacer(1, 12))
            except Exception as e:
                # If table creation fails, add it as preformatted text
                elements.append(Paragraph("Table content (formatting error):", styles['Normal']))
                try:
                    clean_table_content = clean_text_for_pdf(str(part["content"]))
                    elements.append(Preformatted(clean_table_content, styles['Code']))
                except:
                    elements.append(Paragraph("Table could not be displayed", styles['Normal']))
                elements.append(Spacer(1, 12))
    
    # Build the PDF
    try:
        doc.build(elements)
    except Exception as e:
        # If build fails, create a simple PDF with error message and original content
        buffer.seek(0)  # Reset buffer
        buffer.truncate(0)  # Clear buffer
        
        simple_elements = [
            Paragraph("PDF Generation Error", styles['Title']),
            Paragraph(f"Error occurred while generating PDF: {str(e)}", styles['Normal']),
            Spacer(1, 12),
            Paragraph("Original content (as text):", styles['Heading2']),
        ]
        
        # Add original content in chunks to avoid overwhelming the PDF
        try:
            content_preview = str(text)[:5000] + "..." if len(str(text)) > 5000 else str(text)
            clean_content = clean_text_for_pdf(content_preview)
            simple_elements.append(Preformatted(clean_content, styles['Code']))
        except:
            simple_elements.append(Paragraph("Original content could not be displayed", styles['Normal']))
        
        doc.build(simple_elements)
    
    buffer.seek(0)
    return buffer

def generate_docx(text):
    # Generate Word document with proper formatting, including tables
    doc = Document()
    doc.add_heading('AI Response', 0)
    
    # Extract tables and other content
    content_parts = _extract_tables_from_markdown(text)
    
    for part in content_parts:
        if part["type"] == "text":
            # Process non-table text
            paragraphs = part["content"].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    if para.startswith('```') and para.endswith('```'):
                        # Code block - formatted differently
                        code = para.strip('`').strip()
                        p = doc.add_paragraph(code)
                        p.style = 'No Spacing'
                        for run in p.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                        doc.add_paragraph()  # Add space after code
                    else:
                        # Regular paragraph
                        doc.add_paragraph(para)
        
        elif part["type"] == "table":
            # Process table
            table_data = _parse_markdown_table(part["content"])
            if table_data and len(table_data) > 0:
                # Get number of columns from first row
                num_cols = len(table_data[0])
                
                # Create table with appropriate dimensions
                table = doc.add_table(rows=len(table_data), cols=num_cols)
                table.style = 'Table Grid'
                
                # Fill table data
                for i, row in enumerate(table_data):
                    for j, cell_text in enumerate(row):
                        if j < num_cols:  # Ensure we don't exceed column count
                            cell = table.cell(i, j)
                            cell.text = cell_text
                            
                            # Style header row
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                
                doc.add_paragraph()  # Add space after table
    
    # Save to BytesIO object
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def generate_text_file(text):
    # Generate plain text file with original markdown intact
    buffer = BytesIO()
    # Preserve the original text exactly as-is
    buffer.write(text.encode('utf-8'))
    buffer.seek(0)
    return buffer

def generate_excel(text: str) -> BytesIO:
    """
    Generate an Excel file from the given text, with special handling for tables.
    This function fixes the issue with calculating column widths.
    """
    # Create a BytesIO object to store the Excel file
    buffer: BytesIO = BytesIO()

    # Create an Excel writer
    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
        # Extract the different parts of content
        content_parts: list = _extract_tables_from_markdown(text)

        # Track if we've added any content
        has_added_content: bool = False

        # Process tables first - they get priority for proper formatting
        table_count: int = 0
        text_blocks: list = []

        # First pass: extract tables and format them properly
        for part in content_parts:
            if part["type"] == "table":
                # Process tables - each becomes a separate sheet
                table_data: list = _parse_markdown_table(part["content"])
                if table_data and len(table_data) > 0:
                    # Use first row as header
                    headers: list = table_data[0]
                    # Use remaining rows as data
                    data: list = table_data[1:] if len(table_data) > 1 else []

                    # Create DataFrame and clean up data
                    df: pd.DataFrame = pd.DataFrame(data, columns=headers)

                    # Try to convert numeric columns to appropriate types
                    for col in df.columns:
                        try:
                            # Check if column looks numeric (contains mostly digits)
                            numeric_ratio: float = df[col].str.replace('.', '', regex=False).str.isdigit().mean()
                            if not pd.isna(numeric_ratio) and numeric_ratio > 0.5:
                                df[col] = pd.to_numeric(df[col], errors='ignore')
                        except:
                            pass  # Keep as string if conversion fails

                    # Write to Excel sheet with a meaningful name
                    sheet_name: str = f"Table_{table_count+1}"
                    df.to_excel(writer, sheet_name=sheet_name, index=False)

                    # Format the table for better display
                    worksheet = writer.sheets[sheet_name]

                    # Auto-adjust column widths based on content - FIXED CALCULATION
                    for i, col in enumerate(df.columns):
                        # Calculate max length of values in column (convert to string first)
                        col_values_max_len: int = 0
                        try:
                            if not df[col].empty:
                                col_values_max_len = df[col].astype(str).str.len().max()
                                if pd.isna(col_values_max_len):  # Handle NaN result
                                    col_values_max_len = 0
                        except:
                            col_values_max_len = 0
                            
                        # Get length of column name
                        col_name_len: int = len(str(col))
                        
                        # Take maximum of the two lengths + padding
                        max_len: int = max(col_values_max_len, col_name_len) + 2
                        
                        # Excel column index starts at 1
                        col_letter: str = worksheet.cell(row=1, column=i+1).column_letter
                        worksheet.column_dimensions[col_letter].width = min(max_len, 40)  # Cap at 40 characters

                    # Style the header row
                    for cell in worksheet[1]:
                        cell.font = xl.styles.Font(bold=True)
                        cell.fill = xl.styles.PatternFill(start_color="E0E0E0", end_color="E0E0E0", fill_type="solid")

                    table_count += 1
                    has_added_content = True

            elif part["type"] == "text":
                # Collect text blocks for later processing
                if part["content"].strip():
                    text_blocks.append(part["content"].strip())

        # Second pass: process text content if there's any
        if text_blocks:
            # Create a main sheet for text content
            main_sheet_name: str = "Main Content"

            # Process the text to extract meaningful information
            all_text: str = "\n\n".join(text_blocks)

            # Look for list items and key-value pairs
            list_items: list = re.findall(r'(?:^|\n)(?:[\*\-\•\+]|\d+\.)\s+(.+?)(?:\n|$)', all_text)
            key_value_pairs: list = re.findall(r'(?:^|\n)([^:\n]+?):\s*(.+?)(?:\n|$)', all_text)

            # If we found structured data in the text, format it properly
            if key_value_pairs:
                # Create a DataFrame for key-value pairs
                kv_df: pd.DataFrame = pd.DataFrame(key_value_pairs, columns=["Key", "Value"])
                kv_df.to_excel(writer, sheet_name="Key-Value Pairs", index=False)
                has_added_content = True

                # Format the key-value sheet
                kv_sheet = writer.sheets["Key-Value Pairs"]
                # Bold headers
                for cell in kv_sheet[1]:
                    cell.font = xl.styles.Font(bold=True)

                # Auto-adjust column widths - FIXED CALCULATION
                for i, col in enumerate(["Key", "Value"]):
                    # Calculate max length safely
                    col_values_max_len: int = 0
                    try:
                        if not kv_df[col].empty:
                            col_values_max_len = kv_df[col].astype(str).str.len().max()
                            if pd.isna(col_values_max_len):  # Handle NaN result
                                col_values_max_len = 0
                    except:
                        col_values_max_len = 0
                        
                    col_name_len: int = len(col)
                    max_len: int = max(col_values_max_len, col_name_len) + 2
                    
                    col_letter: str = kv_sheet.cell(row=1, column=i+1).column_letter
                    kv_sheet.column_dimensions[col_letter].width = min(max_len, 50)

            if list_items:
                # Create a DataFrame for list items
                list_df: pd.DataFrame = pd.DataFrame({"Items": list_items})
                list_df.to_excel(writer, sheet_name="List Items", index=False)
                has_added_content = True

                # Format the list items sheet
                list_sheet = writer.sheets["List Items"]
                for cell in list_sheet[1]:
                    cell.font = xl.styles.Font(bold=True)

                # Auto-adjust column width - FIXED CALCULATION
                col_values_max_len: int = 0
                try:
                    if not list_df["Items"].empty:
                        col_values_max_len = list_df["Items"].astype(str).str.len().max()
                        if pd.isna(col_values_max_len):  # Handle NaN result
                            col_values_max_len = 0
                except:
                    col_values_max_len = 0
                
                max_len: int = max(col_values_max_len, len("Items")) + 2
                list_sheet.column_dimensions["A"].width = min(max_len, 100)

            # Always include the full text content in a dedicated sheet
            text_df: pd.DataFrame = pd.DataFrame({"Content": [all_text]})
            text_df.to_excel(writer, sheet_name=main_sheet_name, index=False)
            has_added_content = True

            # Format the text content sheet
            text_sheet = writer.sheets[main_sheet_name]
            # Bold header
            for cell in text_sheet[1]:
                cell.font = xl.styles.Font(bold=True)

            # Set column width to a reasonable size
            text_sheet.column_dimensions["A"].width = 100

        # If no content was added (unlikely but possible), create a default sheet
        if not has_added_content:
            df: pd.DataFrame = pd.DataFrame({"Content": [text]})
            df.to_excel(writer, sheet_name="Content", index=False)

            # Format the default sheet
            default_sheet = writer.sheets["Content"]
            for cell in default_sheet[1]:
                cell.font = xl.styles.Font(bold=True)
            default_sheet.column_dimensions["A"].width = 100

    # Reset file pointer to beginning
    buffer.seek(0)
    return buffer

# Process uploaded files
if uploaded_files:
    ui_logger.info(f"Processing {len(uploaded_files)} uploaded files")
    
    for uploaded_file in uploaded_files:
        filename = uploaded_file.name
        file_size = uploaded_file.size if hasattr(uploaded_file, 'size') else 'unknown'
        
        file_logger.info(f"Processing file: {filename} (Size: {file_size} bytes)")
        
        if filename in st.session_state.documents:
            file_logger.debug(f"File {filename} already processed, skipping")
            continue
            
        extracted_text = ""
        
        # Handle different file types with logging
        if filename.lower().endswith(("png", "jpg", "jpeg")):
            file_logger.info(f"Processing image file: {filename}")
            extracted_text = extract_text_from_image(uploaded_file)
        elif filename.lower().endswith(".pdf"):
            file_logger.info(f"Processing PDF file: {filename}")
            extracted_text = extract_text_from_pdf(uploaded_file)
        elif filename.lower().endswith(("xlsx", "xls")):
            file_logger.info(f"Processing Excel file: {filename}")
            extracted_text = extract_text_from_excel(uploaded_file)
        elif filename.lower().endswith(("docx", "doc")):
            file_logger.info(f"Processing Word document: {filename}")
            if filename.lower().endswith(".docx"):
                extracted_text = extract_text_from_docx(uploaded_file)
            else:
                warning_msg = f"Only .docx files are supported. Please convert {filename} to .docx format."
                file_logger.warning(warning_msg)
                extracted_text = warning_msg
        
        if extracted_text.strip():
            st.session_state.documents[filename] = extracted_text
            success_msg = f"File '{filename}' processed. Text extracted successfully."
            file_logger.info(success_msg)
            st.session_state.chat_history.append({
                "type": "system",
                "content": success_msg
            })
        else:
            warning_msg = f"No readable text found in '{filename}'. Try another file."
            file_logger.warning(warning_msg)
            st.warning(warning_msg)
            st.session_state.chat_history.append({
                "type": "system",  
                "content": f"Warning: {warning_msg}"
            })

if st.sidebar.checkbox("Show Logging Status", key="show_logging"):
    st.sidebar.subheader("📊 Logging Status")
    
    # Check if log files exist and show their sizes
    logs_dir = "logs"
    if os.path.exists(logs_dir):
        log_files = ["ocr_app.log", "ocr_app_errors.log"]
        for log_file in log_files:
            log_path = os.path.join(logs_dir, log_file)
            if os.path.exists(log_path):
                size = os.path.getsize(log_path)
                st.sidebar.text(f"{log_file}: {size/1024:.1f} KB")
            else:
                st.sidebar.text(f"{log_file}: Not found")
    
    # Show recent log entries
    if st.sidebar.button("Show Recent Logs"):
        log_path = os.path.join("logs", "ocr_app.log")
        if os.path.exists(log_path):
            try:
                with open(log_path, 'r', encoding='utf-8') as f:
                    lines = f.readlines()
                    recent_lines = lines[-10:]  # Show last 10 lines
                    st.sidebar.text_area("Recent Logs", "".join(recent_lines), height=200)
            except Exception as e:
                st.sidebar.error(f"Could not read log file: {str(e)}")

main_logger.info("Application UI setup completed")

# Document selection UI
if st.session_state.documents:
    st.sidebar.subheader("Available Documents")
    selected_docs = {}
    for doc_name in st.session_state.documents.keys():
        selected = st.sidebar.checkbox(f"{doc_name}", key=f"select_{doc_name}")
        selected_docs[doc_name] = selected
    st.session_state.selected_docs = selected_docs

    # Retrieve compare mode checkbox value without assigning it to session_state
    compare_mode = st.sidebar.checkbox("Compare selected documents", key="compare_mode")

    st.subheader("Selected Document Contents:")
    any_selected = any(selected_docs.values())
    if any_selected:
        for doc_name, is_selected in selected_docs.items():
            if is_selected:
                with st.expander(f"📄 {doc_name}"):
                    st.text_area(f"Content of {doc_name}",
                                 st.session_state.documents[doc_name],
                                 height=150,
                                 key=f"content_{doc_name}")
    else:
        st.info("Select one or more documents from the sidebar to view their contents.")
else:
    compare_mode = False  # Ensure compare_mode is defined when no documents exist

# User input - only show when documents are available
if st.session_state.documents:
    user_prompt = st.chat_input("Enter your prompt about the documents:")

    if user_prompt:
        logging.info(f"User prompt submitted: {user_prompt}")
        # Prevent system messages or responses being treated as user input
        if ":" in user_prompt and any(prefix in user_prompt for prefix in ["👤 User:", "🤖 AI:", "🖥️ System:"]):
            st.warning("Please enter your own prompt without copying previous messages.")
        else:
            st.session_state.original_prompt = user_prompt
            # MODIFIED: No automatic transformation of the prompt
            st.session_state.transformed_prompt = user_prompt  # Use original prompt as-is
            st.session_state.needs_processing = True
            st.rerun()
else:
    # Show message when no documents are available
    st.info("📤 Upload a document first to start chatting with AI.")
    # Disable the chat input when no documents exist
    st.text_input("Enter your prompt (upload documents first):", disabled=True)

# Chat History Display
st.subheader("💬 Chat History")
for i, entry in enumerate(st.session_state.chat_history):
    if entry["type"] == "user":
        st.markdown(f"👤 **User**: {entry['content']}", unsafe_allow_html=True)
    elif entry["type"] == "system":
        st.markdown(f"🖥️ **System**: {entry['content']}", unsafe_allow_html=True)
    elif entry["type"] == "ai":
        st.markdown(f"🤖 **AI**: {entry['content']}", unsafe_allow_html=True)
    if i < len(st.session_state.chat_history) - 1:
        st.markdown("---")

# Document + Prompt Processing
if st.session_state.needs_processing and st.session_state.original_prompt and st.session_state.documents:
    selected_doc_names = [name for name, selected in st.session_state.selected_docs.items() if selected]
    if selected_doc_names:
        st.session_state.chat_history.append({
            "type": "user",
            "content": f"Prompt: {st.session_state.original_prompt}"
        })
        if compare_mode and len(selected_doc_names) > 1:
            combined = []
            for i, doc_name in enumerate(selected_doc_names):
                combined.append(f"Document {i+1} ({doc_name}):\n{st.session_state.documents[doc_name]}")
            full_text = "\n\n---\n\n".join(combined)
            response = generate_chat_response(full_text, st.session_state.original_prompt, is_comparison=True)
            st.session_state.chat_history.append({
                "type": "system",
                "content": f"Comparing documents: {', '.join(selected_doc_names)}"
            })
            st.session_state.chat_history.append({
                "type": "ai",
                "content": response
            })
        else:
            for doc_name in selected_doc_names:
                doc_text = st.session_state.documents[doc_name]
                st.session_state.chat_history.append({
                    "type": "system",
                    "content": f"Processing document: {doc_name}"
                })
                response = generate_chat_response(doc_text, st.session_state.original_prompt)
                st.session_state.chat_history.append({
                    "type": "ai",
                    "content": response
                })
    else:
        st.session_state.chat_history.append({
            "type": "user",
            "content": st.session_state.original_prompt
        })
        if st.session_state.documents:
            st.session_state.chat_history.append({
                "type": "system",
                "content": "Please select one or more documents from the sidebar to process your query."
            })
        else:
            response = generate_chat_response("", st.session_state.original_prompt)
            st.session_state.chat_history.append({
                "type": "ai",
                "content": response
            })

    st.session_state.transformed_prompt = ""
    st.session_state.original_prompt = ""
    st.session_state.needs_processing = False
    st.rerun()

# Download AI responses in multiple formats
ai_responses = [entry["content"] for entry in st.session_state.chat_history if entry["type"] == "ai"]
if ai_responses:
    st.write("🤖 AI Generated Responses:")
    
    # Create a selectbox to choose which response to download
    response_options = ["All responses combined"]
    
    # Create better descriptions for individual responses
    for i, response in enumerate(ai_responses):
        # Get a preview of the response
        preview = response[:50] + "..." if len(response) > 50 else response
        response_options.append(f"Response #{i+1}: {preview}")
    
    selected_option = st.selectbox(
        "Select which response to download:",
        response_options,
        key="response_selector"
    )
    
    # Determine which response content to use
    if selected_option == "All responses combined":
        download_content = "\n\n---\n\n".join(ai_responses)
        file_prefix = "all_responses"
    else:
        # More robust extraction of response number
        import re
        response_pattern = re.search(r"Response #(\d+)", selected_option)
        if response_pattern:
            response_num = int(response_pattern.group(1)) - 1
            if 0 <= response_num < len(ai_responses):
                download_content = ai_responses[response_num]
                file_prefix = f"response_{response_num+1}"
            else:
                st.error(f"Invalid response number: {response_num+1}")
                download_content = "\n\n---\n\n".join(ai_responses)
                file_prefix = "all_responses"
        else:
            st.error("Could not parse response number, downloading all responses")
            download_content = "\n\n---\n\n".join(ai_responses)
            file_prefix = "all_responses"
    
    # Create columns for download buttons
    col1, col2, col3, col4 = st.columns(4)  # Added a fourth column for Excel
    
    # PDF download
    with col1:
        pdf_buffer = generate_pdf(download_content)
        st.download_button(
            label="📄 Download as PDF",
            data=pdf_buffer,
            file_name=f"{file_prefix}.pdf",
            mime="application/pdf",
            key=f"download_{file_prefix}_pdf"
        )
    
    # Word document download
    with col2:
        docx_buffer = generate_docx(download_content)
        st.download_button(
            label="📝 Download as Word",
            data=docx_buffer,
            file_name=f"{file_prefix}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"download_{file_prefix}_docx"
        )
    
    # Plain text download
    with col3:
        txt_buffer = generate_text_file(download_content)
        st.download_button(
            label="📋 Download as Text",
            data=txt_buffer,
            file_name=f"{file_prefix}.txt",
            mime="text/plain",
            key=f"download_{file_prefix}_txt"
        )
    
    # Excel download - NEW
    with col4:
        excel_buffer = generate_excel(download_content)
        st.download_button(
            label="📊 Download as Excel",
            data=excel_buffer,
            file_name=f"{file_prefix}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"download_{file_prefix}_excel"
        )

with st.sidebar.expander("📘 How to Use This Tool", expanded=False):
    st.markdown("""
1. **Upload Your Files**  
   - Use the menu on the left to upload your files.  
   - Supported file types include:  
     - Images (e.g., scanned documents)  
     - PDF files  
     - Word documents  
     - Excel spreadsheets  

2. **View Extracted Text**  
   - Once files are uploaded, select them using the checkboxes.  
   - The extracted text from each file will be displayed on the screen.

3. **Ask Questions About the Content**  
   - Use the input box at the bottom to ask questions.  
   - Example queries include:  
     - "Summarize this document"  
     - "What are the key points?"  
     - "Compare the two documents"

4. **Compare Documents**  
   - To compare, simply check the **"Compare selected documents"** option.

5. **View Answers**  
   - The tool will respond with summaries or relevant answers based on the file content.

6. **Download Results**  
   - Scroll down to download your results.  
   - Available formats: PDF, Word, plain text, or spreadsheet.

7. **Clear the Chat**  
   - Use the **"Clear Chat"** button to reset and start a new session.

8. **Optimize Your Prompts**  
   - For more accurate and useful results, write clear and meaningful prompts.  
   - Use the **"Optimize Prompt"** section in the sidebar to rephrase your question.  
   - Copy the optimized prompt and paste it into the main input box, then press **Enter** to see improved output.
    """)